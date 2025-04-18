import gradio as gr
import whisper
from docx import Document
from datetime import datetime
from transformers import pipeline
from deep_translator import GoogleTranslator
from gtts import gTTS
import os
from langdetect import detect

# Initialize a list to hold pre-added members and their transcriptions
members = []
transcriptions = []  # Changed from dict to list to maintain order


model = whisper.load_model("small")



# Initialize the summarization model
summarizer = pipeline("summarization")

# Function to detect and translate Hindi to Urdu
def detect_and_translate_to_urdu(text):
    if detect(text) == 'hi':  # Detect Hindi
        return GoogleTranslator(source='auto', target='ur').translate(text)  # Translate to Urdu
    return text


def transcribe_audio(audio):
    if not audio:
        return "No audio file provided.", ""

    try:
        # Transcribe directly from file path
        result = model.transcribe(audio)
        transcription = result["text"]

        # Detect and translate Hindi to Urdu
        transcription_in_urdu = detect_and_translate_to_urdu(transcription)

        return transcription_in_urdu, ""

    except Exception as e:
        return f"Error: {e}", ""




def summarize_transcription(english_translation):
    if not english_translation:
        return "No translated text available for summarization."

    summary = summarizer(english_translation, max_length=150, min_length=50, do_sample=False)[0]["summary_text"]
    return summary




def translate_urdu_to_english(transcription_in_urdu):
    if not transcription_in_urdu:
        return "No transcription available to translate."

    translation = GoogleTranslator(source='ur', target='en').translate(transcription_in_urdu)
    return translation




# Function to add members
def add_member(new_member, members_list):
    if new_member and new_member not in members_list:
        members_list.append(new_member)  # Append only for this session
    return "\n".join(members_list), "", members_list  # Return updated list


def display_minutes(meeting_title, date, venue, requested_by, agenda, action_items, next_meeting, additional_notes, members_list, summary_output):
    minutes = f"{meeting_title}\n\n"
    minutes += f"Date and Time: {date}\n"
    minutes += f"Venue: {venue}\n"
    minutes += f"Requested by: {requested_by}\n\n"

    minutes += "Meeting Attendees:\n"
    if members_list:
        for member in members_list:
            minutes += f"  - {member}\n"
    else:
        minutes += "  No attendees listed.\n"

    minutes += "\nAgenda:\n"
    minutes += f"{agenda}\n\n"

    minutes += "Discussion:\n"
    if summary_output:
        minutes += f"  {summary_output}\n"
    else:
        minutes += "  No summary available.\n"

    minutes += "\nAction Items:\n"
    minutes += f"{action_items}\n\n"

    minutes += "Next Meeting:\n"
    minutes += f"{next_meeting}\n\n"

    minutes += "Additional Notes:\n"
    minutes += f"{additional_notes}\n\n"

    return minutes


"""
# Function to save meeting minutes as a Word document
def save_minutes(minutes, file_name="Meeting_Minutes.docx"):
    try:
        doc = Document()
        doc.add_heading("Meeting Minutes", 0)

        for line in minutes.split("\n"):
            if line.strip() == "":
                doc.add_paragraph()
            else:
                doc.add_paragraph(line)

        doc.save(file_name)
        return file_name
    except Exception as e:
        return f"An error occurred while saving the document: {e}"

"""
import gradio as gr
from docx import Document
from docx.shared import Pt

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

def save_minutes(meeting_title, date, venue, requested_by, agenda, action_items, next_meeting, additional_notes, members_list, summary_output):
    try:
        doc = Document()

        # --- Title Formatting (Centered & Bold) ---
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(meeting_title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("\n")  # Spacer

        # --- Table for Meeting Details ---
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        # Add meeting details in table format
        details = [("Date:", date), ("Venue:", venue), ("Requested by:", requested_by)]
        
        for i, (label, value) in enumerate(details):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value

        # --- Meeting Attendees ---
        attendees_text = "\n".join([f"- {member}" for member in members_list]) if members_list else "No attendees listed."
        table.cell(3, 0).text = "Meeting Attendees:"
        table.cell(3, 1).text = attendees_text

        doc.add_paragraph("\n")  # Spacer

        # --- Agenda ---
        doc.add_paragraph("Agenda:", style="Heading 1")
        doc.add_paragraph(agenda)

        # --- Discussion / Summary ---
        doc.add_paragraph("\nDiscussion:", style="Heading 1")
        doc.add_paragraph(summary_output if summary_output else "No summary available.")

        # --- Action Items ---
        doc.add_paragraph("\nAction Items:", style="Heading 1")
        doc.add_paragraph(action_items)

        # --- Next Meeting ---
        doc.add_paragraph("\nNext Meeting:", style="Heading 1")
        doc.add_paragraph(next_meeting)

        # --- Additional Notes ---
        doc.add_paragraph("\nAdditional Notes:", style="Heading 1")
        doc.add_paragraph(additional_notes)

        # Save the document
        file_name = "Meeting_Minutes.docx"
        doc.save(file_name)
        return file_name

    except Exception as e:
        return f"An error occurred while saving the document: {e}"












# Function to convert text to speech (TTS) using the transcribed text
def text_to_audio_from_transcription(transcribed_text):
    if not transcribed_text:
        return "No text provided for conversion."

    try:
        tts = gTTS(text=transcribed_text, lang="ur")
        file_path = "meeting_audio.mp3"
        tts.save(file_path)
        return file_path  # Return the file path for playback and download
    except Exception as e:
        return f"Error in text-to-audio conversion: {e}"


import shutil

# Function to save recorded audio
def save_audio(audio_path):
    if not audio_path:
        return "No audio recorded.", None

    # Define a filename for the saved audio
    saved_audio_path = "saved_recorded_audio.wav"
    shutil.copy(audio_path, saved_audio_path)  # Copy to make it downloadable

    return f"Audio saved successfully: {saved_audio_path}", saved_audio_path


# Create the Gradio interface
with gr.Blocks(theme=gr.themes.Soft(primary_hue="slate", secondary_hue="zinc")) as demo:

    gr.HTML("<h1 style='text-align: center; color: black;'>Conference Management System</h1>")
    gr.HTML("<p style='text-align: center; color: black;'>This system allows you to upload and record audios plus transcribe, translate, summarize, and generate meeting minutes.</p>")

    
    with gr.Row():
        meeting_title_input = gr.Textbox(label="Meeting Title", placeholder="Enter the meeting title...")
        date_input = gr.Textbox(label="Select Meeting Date", value=datetime.now().strftime('%Y-%m-%d'), placeholder="YYYY-MM-DD")
        time_input = gr.Textbox(label="Time (e.g., 02:00pm - 04:00pm)")
        venue_input = gr.Textbox(label="Venue", placeholder="Enter the venue...")
        requested_by_input = gr.Textbox(label="Requested by", placeholder="Enter the name of the person who requested the meeting...")
        apologies_input = gr.Textbox(label="Apologies (comma-separated)")

    with gr.Row():
        members_display = gr.Textbox(label="List of Members", value="", interactive=False, lines=5)
        new_member_input = gr.Textbox(label="Add New Member")
        add_member_button = gr.Button("Add Member")
        audio_input = gr.Audio(label="Upload/Record Audio", type="filepath")  # Allows upload & real-time recording
        savew_button = gr.Button("Save Recorded Audio")
        status_output = gr.Textbox(label="Status", interactive=False)
        download_audio = gr.File(label="Download Recorded Audio")


    with gr.Row():
        agenda_input = gr.Textbox(label="Agenda", placeholder="List the agenda points here...")
        action_items_input = gr.Textbox(label="Action Items", placeholder="List the action items here...")
        next_meeting_input = gr.Textbox(label="Next Meeting", placeholder="Details for the next meeting...")
        additional_notes_input = gr.Textbox(label="Additional Notes", placeholder="Any additional notes...")


    with gr.Row():
        transcribe_button = gr.Button("Transcribe Audio")
        translate_urdu_to_english_button = gr.Button("Translate Urdu to English")
        summarize_button = gr.Button("Display Key Points")
        display_button = gr.Button("Display Meeting Minutes")
        save_button = gr.Button("Save Meeting Minutes")
        text_to_audio_button = gr.Button("Convert Text to Audio")

    members_state = gr.State([])  # This ensures a fresh list every new Gradio launch
    # transcriptions_state = gr.State([])  # Stores transcriptions per session

    transcribed_text_output = gr.Textbox(label="Transcribed Text", interactive=True)
    urdu_to_english_output = gr.Textbox(label="Translated Urdu to English Text", interactive=True)
    summary_output = gr.Textbox(label="Key Points", interactive=True)
    minutes_output = gr.Textbox(label="Meeting Minutes", interactive=True)
    file_output = gr.File(label="Download Meeting Minutes Document")
    audio_output = gr.Audio(label="Generated Speech", type="filepath")

    add_member_button.click(fn=add_member, inputs=[new_member_input, members_state],outputs=[members_display, new_member_input, members_state])
    transcribe_button.click(fn=transcribe_audio, inputs=[audio_input], outputs=[transcribed_text_output, minutes_output])

    translate_urdu_to_english_button.click(
      fn=translate_urdu_to_english,
      inputs=[transcribed_text_output],  # Takes Urdu transcription
      outputs=[urdu_to_english_output]   # Outputs English translation
)

    summarize_button.click(
      fn=summarize_transcription,
      inputs=[urdu_to_english_output],  # Takes translated text as input
      outputs=[summary_output]          # Outputs summarized text
)



    display_button.click(
      fn=display_minutes,
      inputs=[meeting_title_input, date_input, venue_input, requested_by_input,
            agenda_input, action_items_input, next_meeting_input, additional_notes_input,
            members_state, summary_output],  # Uses summary as input
    outputs=[minutes_output]
)

    save_button.click(
    fn=save_minutes,
    inputs=[
        meeting_title_input,
        date_input,
        venue_input,
        requested_by_input,
        agenda_input,
        action_items_input,
        next_meeting_input,
        additional_notes_input,
        members_state,  # Assuming this holds members_list
        summary_output      # Assuming this holds the summary_output
    ],
    outputs=[file_output]
)




    savew_button.click(fn=save_audio, inputs=audio_input, outputs=[status_output, download_audio])
      # Update the text_to_audio_button to use the transcribed text from the textbox
    text_to_audio_button.click(
      fn=text_to_audio_from_transcription,
      inputs=[urdu_to_english_output],  # Use the English translation as input
      outputs=[audio_output]  # Output the generated speech file
)



demo.launch(share=True)

